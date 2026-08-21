"""Contract Review service: extraction, optional OCR, and structured LLM analysis."""
from __future__ import annotations

import hashlib
import io
import json
import os
import re
import tempfile
import threading
import time
import uuid
from pathlib import Path
from typing import Any

from docx import Document
from openai import OpenAI
from pypdf import PdfReader

try:
    from google import genai
    from google.genai import types as genai_types
except ImportError:
    genai = None
    genai_types = None

try:
    from PIL import Image, ImageOps
except ImportError:
    Image = None
    ImageOps = None

MAX_REVIEW_BYTES = int(os.getenv("REVIEW_MAX_FILE_BYTES", str(20 * 1024 * 1024)))
MAX_REVIEW_CHARS = int(os.getenv("REVIEW_MAX_TEXT_CHARS", "120000"))
OLLAMA_BASE_URL = os.getenv("OLLAMA_BASE_URL", "http://localhost:11434/v1")
REVIEW_MODEL = os.getenv("REVIEW_LLM_MODEL", os.getenv("LLM_MODEL", "qwen3:14b"))
GEMINI_KEY = os.getenv("GEMINI_KEY") or os.getenv("GEMINI_API_KEY")
GEMINI_OCR_MODEL = os.getenv("GEMINI_OCR_MODEL", "gemma-4-31b-it")
OCR_CACHE_DIR = Path(os.getenv("REVIEW_OCR_CACHE_DIR", str(Path(__file__).resolve().parents[1] / "ocr_cache")))
OCR_CACHE_DIR.mkdir(parents=True, exist_ok=True)

_jobs: dict[str, dict[str, Any]] = {}
_jobs_lock = threading.RLock()


def get_job(job_id: str) -> dict[str, Any] | None:
    with _jobs_lock:
        job = _jobs.get(job_id)
        return dict(job) if job else None


def create_job(*, file_path: str, filename: str, content_type: str, session_id: str,
               user_id: str, firm_id: str, on_completed=None) -> str:
    job_id = str(uuid.uuid4())
    with _jobs_lock:
        _jobs[job_id] = {
            "job_id": job_id,
            "status": "queued",
            "progress": 0.0,
            "stage": "تم استلام الملف",
            "filename": filename,
            "content_type": content_type,
            "session_id": session_id,
            "user_id": user_id,
            "firm_id": firm_id,
            "file_path": file_path,
            "source_text": None,
            "result": None,
            "error": None,
            "on_completed": on_completed,
        }
    threading.Thread(target=_run_job, args=(job_id,), daemon=True).start()
    return job_id


def job_belongs_to_user(job: dict[str, Any], user_id: str, firm_id: str) -> bool:
    return job.get("user_id") == user_id and job.get("firm_id") == firm_id


def _set_job(job_id: str, **patch: Any) -> None:
    with _jobs_lock:
        if job_id in _jobs:
            _jobs[job_id].update(patch)


def extract_document_text(path: str, filename: str, content_type: str) -> tuple[str, str]:
    """Return (text, extraction_method). OCR is explicit and never silently faked."""
    suffix = Path(filename).suffix.lower()
    if suffix in {".txt", ".md", ".csv"} or content_type.startswith("text/"):
        text = Path(path).read_text(encoding="utf-8", errors="replace")
        return _limit_text(text), "text"

    if suffix == ".pdf" or content_type == "application/pdf":
        reader = PdfReader(path)
        text = "\n\n".join((page.extract_text() or "") for page in reader.pages).strip()
        if text:
            return _limit_text(text), "pdf-text"
        return _gemini_ocr_file(path, filename, "application/pdf")

    if suffix == ".docx" or content_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
        doc = Document(path)
        blocks = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                if any(cells):
                    blocks.append(" | ".join(cells))
        text = "\n\n".join(blocks).strip()
        if not text:
            raise ValueError("ملف Word لا يحتوي على نص قابل للاستخراج")
        return _limit_text(text), "docx-text"

    if suffix in {".png", ".jpg", ".jpeg", ".tif", ".tiff", ".webp"} or content_type.startswith("image/"):
        return _gemini_ocr_file(path, filename, content_type or "application/octet-stream")

    raise ValueError("نوع الملف غير مدعوم. استخدمي PDF أو DOCX أو TXT أو صورة.")


def _gemini_ocr_file(path: str, filename: str, content_type: str) -> tuple[str, str]:
    """Extract scanned contract text with the Gemini Vision flow from contractgeneration."""
    if genai is None or genai_types is None:
        raise RuntimeError("حزمة google-genai غير مثبتة؛ لا يمكن تشغيل Gemini OCR")
    if not GEMINI_KEY:
        raise RuntimeError("GEMINI_KEY أو GEMINI_API_KEY غير موجود في backend/.env")

    raw = Path(path).read_bytes()
    cache_key = hashlib.sha256(raw + GEMINI_OCR_MODEL.encode("utf-8")).hexdigest()
    cache_file = OCR_CACHE_DIR / f"{cache_key}.txt"
    if cache_file.exists():
        cached = _limit_text(cache_file.read_text(encoding="utf-8", errors="replace"))
        if cached:
            return cached, "gemini-ocr-cache"

    payload = raw
    mime_type = content_type
    if content_type.startswith("image/") and Image is not None and ImageOps is not None:
        image = Image.open(io.BytesIO(raw)).convert("L")
        image = ImageOps.autocontrast(image)
        threshold = int(os.getenv("GEMINI_OCR_THRESHOLD", "180"))
        image = image.point(lambda pixel: 255 if pixel > threshold else 0)
        buffer = io.BytesIO()
        image.save(buffer, format="PNG")
        payload = buffer.getvalue()
        mime_type = "image/png"

    prompt = """أنت خبير OCR متخصص في العقود العربية. استخرج كل النص بدقة متناهية سطراً بسطر.
القواعد:
- النص المطبوع: انقله كما هو.
- النص اليدوي في الفراغات: [يدوي: النص].
- فراغ فارغ: [فارغ].
- كلمة غير واضحة: [غير واضح].
- حافظ على ترتيب وتنسيق السطور.
- لا تضف أي شرح أو تلخيص أو كلام من عندك؛ أعد النص المستخرج فقط.
"""
    client = genai.Client(api_key=GEMINI_KEY)
    last_error: Exception | None = None
    for attempt in range(3):
        try:
            response = client.models.generate_content(
                model=GEMINI_OCR_MODEL,
                contents=[prompt, genai_types.Part.from_bytes(data=payload, mime_type=mime_type)],
            )
            text = _limit_text(getattr(response, "text", "") or "")
            if not text:
                raise RuntimeError("Gemini OCR لم يعطِ نصًا واضحًا")
            cache_file.write_text(text, encoding="utf-8")
            return text, "gemini-ocr"
        except Exception as exc:
            last_error = exc
            if attempt < 2:
                time.sleep((attempt + 1) * 3)
    raise RuntimeError(f"فشل Gemini OCR للملف {filename}: {last_error}")


def _limit_text(text: str) -> str:
    normalized = re.sub(r"\n{3,}", "\n\n", text).strip()
    if len(normalized) > MAX_REVIEW_CHARS:
        return normalized[:MAX_REVIEW_CHARS] + "\n[تم اختصار النص بسبب حد الطول]"
    return normalized


def analyze_contract(text: str) -> dict[str, Any]:
    prompt = f"""أنت مراجع عقود قانوني. حلّل النص التالي تحليلاً منظماً، ولا تخترع مادة قانونية أو واقعة غير موجودة.
أعد JSON صالحاً فقط بالحقول التالية:
{{
  "title": "string",
  "summary": "string",
  "overall_risk": "safe|review|risk",
  "overall_score": 0,
  "clauses": [{{
    "number": "string",
    "title": "string",
    "excerpt": "string",
    "status": "safe|review|risk",
    "risk_score": 0,
    "reason": "string",
    "legal_ref": "string|null",
    "legal_basis": "string|null",
    "recommendation": "string"
  }}],
  "recommendations": ["string"],
  "disclaimer": "هذه مراجعة آلية أولية وليست رأياً قانونياً نهائياً."
}}
قواعد مهمة:
- صنّف المخاطر بناءً على النص فقط.
- إذا لم يوجد مرجع قانوني مؤكد، استخدم null ولا تخترع رقم مادة.
- استخرج البنود الموجودة فعلاً حتى لو كانت صياغتها ناقصة.
- overall_score وrisk_score أرقام من 0 إلى 100.

نص العقد:
{text}
"""
    client = OpenAI(api_key="ollama", base_url=OLLAMA_BASE_URL)
    response = client.chat.completions.create(
        model=REVIEW_MODEL,
        messages=[
            {"role": "system", "content": "أخرج JSON فقط بلا Markdown."},
            {"role": "user", "content": prompt},
        ],
        temperature=0.1,
        max_tokens=int(os.getenv("REVIEW_MAX_TOKENS", "5000")),
    )
    raw = (response.choices[0].message.content or "").strip()
    data = _parse_json(raw)
    return _validate_result(data)


def _parse_json(raw: str) -> dict[str, Any]:
    cleaned = re.sub(r"<think[\s\S]*?</think>", "", raw, flags=re.IGNORECASE).strip()
    cleaned = re.sub(r"^```(?:json)?\s*|\s*```$", "", cleaned, flags=re.IGNORECASE).strip()
    try:
        parsed = json.loads(cleaned)
    except json.JSONDecodeError:
        match = re.search(r"\{[\s\S]*\}", cleaned)
        if not match:
            raise RuntimeError("الموديل أعاد نتيجة غير منظمة لمراجعة العقد")
        parsed = json.loads(match.group(0))
    if not isinstance(parsed, dict):
        raise RuntimeError("نتيجة مراجعة العقد ليست كائن JSON")
    return parsed


def _validate_result(data: dict[str, Any]) -> dict[str, Any]:
    allowed = {"safe", "review", "risk"}
    clauses = []
    for index, raw_clause in enumerate(data.get("clauses") or [], start=1):
        if not isinstance(raw_clause, dict):
            continue
        status = raw_clause.get("status") if raw_clause.get("status") in allowed else "review"
        clauses.append({
            "number": str(raw_clause.get("number") or index),
            "title": str(raw_clause.get("title") or f"البند {index}"),
            "excerpt": str(raw_clause.get("excerpt") or ""),
            "status": status,
            "risk_score": max(0, min(100, int(raw_clause.get("risk_score") or 0))),
            "reason": str(raw_clause.get("reason") or ""),
            "legal_ref": raw_clause.get("legal_ref"),
            "legal_basis": raw_clause.get("legal_basis"),
            "recommendation": str(raw_clause.get("recommendation") or ""),
        })
    overall = data.get("overall_risk") if data.get("overall_risk") in allowed else "review"
    return {
        "title": str(data.get("title") or "مراجعة عقد"),
        "summary": str(data.get("summary") or ""),
        "overall_risk": overall,
        "overall_score": max(0, min(100, int(data.get("overall_score") or 0))),
        "clauses": clauses,
        "recommendations": [str(x) for x in (data.get("recommendations") or []) if x],
        "disclaimer": str(data.get("disclaimer") or "هذه مراجعة آلية أولية وليست رأياً قانونياً نهائياً."),
    }


def _run_job(job_id: str) -> None:
    with _jobs_lock:
        job = _jobs.get(job_id)
        if not job:
            return
        path = job["file_path"]
        filename = job["filename"]
        content_type = job["content_type"]
    try:
        _set_job(job_id, status="processing", progress=0.1, stage="استخراج النص")
        text, method = extract_document_text(path, filename, content_type)
        _set_job(job_id, progress=0.35, stage="تم استخراج النص", source_text=text, extraction_method=method)
        _set_job(job_id, progress=0.55, stage="تحليل البنود والمخاطر")
        result = analyze_contract(text)
        result["extraction_method"] = method
        _set_job(job_id, status="completed", progress=1.0, stage="اكتملت مراجعة العقد", result=result)
        callback = job.get("on_completed")
        if callback:
            try:
                callback(job_id, text, method, result)
            except Exception as persistence_error:
                _set_job(job_id, persistence_error=str(persistence_error))
    except Exception as exc:
        _set_job(job_id, status="failed", progress=1.0, stage="فشلت المراجعة", error=str(exc))
    finally:
        try:
            os.unlink(path)
        except OSError:
            pass
