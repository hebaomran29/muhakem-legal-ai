import warnings
warnings.filterwarnings("ignore")
import os
os.environ["TRANSFORMERS_VERBOSITY"] = "error"
os.environ["TOKENIZERS_PARALLELISM"] = "false"

import re
import json
from qdrant_client import QdrantClient
from sentence_transformers import SentenceTransformer
from openai import OpenAI  # هنستخدمه للاتصال بـ Ollama (OpenAI-compatible endpoint)
from pypdf import PdfReader
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from dotenv import load_dotenv

load_dotenv()  # بيقرأ ملف .env من نفس مجلد المشروع (لازم يكون فيه QDRANT_URL و QDRANT_KEY)

# ──────────────────────────────────────────────────────────────────
# إعدادات الاتصال والمسارات
# ──────────────────────────────────────────────────────────────────
QDRANT_URL = os.environ.get("QDRANT_URL")
QDRANT_KEY = os.environ.get("QDRANT_KEY")

OLLAMA_MODEL = "qwen3:14b"
OLLAMA_BASE_URL = "http://localhost:11434/v1"
CLAUSE_TYPES_PATH = os.path.join(os.path.dirname(os.path.abspath(__file__)),
                                    "muhakkim_specific_clauses.json")

# توليد بند واحد بيحتاج توكينز أقل بكتير من توليد العقد كله دفعة واحدة
CLAUSE_MAX_TOKENS = 2500
INTRO_MAX_TOKENS = 1200

# ──────────────────────────────────────────────────────────────────
# إعدادات الـ RAG
# ──────────────────────────────────────────────────────────────────
COLLECTION_LAWS = "laws_only"        # القوانين والتشريعات (المصدر الملزم)
COLLECTION_CONTRACTS = "legal_contracts"        # عقود مرجعية / سوابق (استرشادية فقط)
EMBED_MODEL_CONTRACTS = "intfloat/multilingual-e5-base"
EMBED_MODEL_LAWS = "mohamed2811/Muffakir_Embedding_V2"
LAWS_SCORE_THRESHOLD = 0.72

# ──────────────────────────────────────────────────────────────────
# خريطة الكلمات المفتاحية لتحديد نوع العقد
# ──────────────────────────────────────────────────────────────────
CONTRACT_TYPE_KEYWORDS = {
    "sale": ["بيع", "شراء", "مبيع", "بائع", "مشتري", "بيع وشراء"],
    "lease": ["ايجار", "إيجار", "استئجار", "مؤجر", "مستأجر", "تأجير"],
    "contracting": ["مقاولة", "مقاول", "مقاولات", "تنفيذ أعمال", "بناء", "إنشاءات"],
    "partnership": ["شراكة", "شركة", "شريك", "تأسيس شركة", "حصص"],
    "deposit": ["وديعة", "إيداع", "ايداع", "استيداع", "مودع"],
    "guarantee": ["كفالة", "ضمان", "كفيل", "ضامن"],
    "employment_contract": ["عمل", "توظيف", "موظف", "عامل", "وظيفة", "تشغيل"],
    "agency": ["وكالة", "وكيل", "موكل", "تفويض"],
}

# ──────────────────────────────────────────────────────────────────
# السيستم برومبت الخاص بصياغة *بند واحد فقط* (مش العقد كله)
#
# ✏️ تعديل 1 (الأهم): الأولوية اتقلبت بالكامل.
# قبل كده كان مكتوب "لو فيه أرقام في وصف الالتزام التزمي بيها حرفياً" —
# وده اللي خلّى الموديل يعتمد على أرقام الـ JSON (اللي ممكن تتلخبط زي
# حالة 180/90 يوم) بدل ما يستنى النص القانوني الحقيقي من RAG.
# دلوقتي: description بقى مجرد "موضوع البند" (توجيه)، والرقم الملزم
# الوحيد اللي الموديل مسموحله يستخدمه هو اللي جاي في "المرجع القانوني"
# (لو موجود). لو مفيش مرجع قانوني، الموديل يكتب صياغة عامة بلا رقم.
# ──────────────────────────────────────────────────────────────────
CLAUSE_SYSTEM_PROMPT = """أنت مستشار قانوني مصري متخصص في صياغة العقود بالعربية الفصحى الرسمية.

مهمتك محدودة جداً: صياغة **فقرة قانونية واحدة فقط** (بند واحد) تعبّر عن الالتزام
التالي بأسلوب تعاقدي ملزم، دقيق، ومفصّل (بحد أدنى 100-120 كلمة).

# قواعد صارمة:
- ممنوع كتابة عنوان البند أو رقمه أو كلمة "البند" — أنا (الكود) اللي هحطها، انت تكتب نص الفقرة بس.
- ممنوع أي رموز Markdown (** أو # أو _).
- ممنوع الترحيب أو أي مقدمة، ابدأي بصلب الفقرة القانونية مباشرة.
- **أهم قاعدة (مصدر الأرقام):** "المرجع القانوني" بالأسفل (لو موجود) هو المصدر
  الوحيد الملزم لأي رقم أو نسبة أو مدة زمنية. لو فيه رقم صريح جوه المرجع
  القانوني، استخدميه حرفياً. **ممنوع تماماً** إنك تاخدي أي رقم من "وصف
  الالتزام" لوحده لو مش مؤكَّد بنفس الرقم داخل المرجع القانوني.
- لو "المرجع القانوني" مش موجود بالأسفل، اكتبي الالتزام بصياغة عامة دقيقة
  قانونياً **بدون** اختراع أي رقم أو نسبة أو مدة من عندك. وصف الالتزام هنا
  بيوضحلك موضوع البند فقط، مش مصدر أرقام يُعتمد عليه.
- ممنوع الاختراع أو الإضافة من خارج البيانات المُعطاة لك.
- لا تذكري رقم المادة القانونية في نص البند، اكتفي بصياغة المضمون القانوني.

# بيانات هذا البند تحديداً:
نوع العقد: {contract_type_ar}
عنوان البند (الموضوع): {clause_title}
وصف الالتزام (توجيه عام للموضوع فقط، مش مصدر أرقام): {clause_description}
{laws_context_block}
اكتبي فقرة الالتزام الآن مباشرة بدون أي عنوان أو ترقيم."""

# ──────────────────────────────────────────────────────────────────
# تحميل موارد الـ RAG
# ──────────────────────────────────────────────────────────────────
def load_rag_resources() -> dict:
    resources = {
        "qdrant_client": None,
        "embed_model": None,
        "embed_model_laws": None,
        "contracts_ready": False,
        "laws_ready": False,
    }

    if not QDRANT_URL or not QDRANT_KEY:
        print("❌ QDRANT_URL أو QDRANT_KEY مش متعرّفين — تأكدي من ملف .env")
        return resources

    try:
        qdrant = QdrantClient(
            url=QDRANT_URL,
            api_key=QDRANT_KEY,
            port=443,
            https=True,
            timeout=30,
            check_compatibility=False
        )
        resources["qdrant_client"] = qdrant

        collections = qdrant.get_collections().collections
        collection_names = [c.name for c in collections]

        resources["contracts_ready"] = COLLECTION_CONTRACTS in collection_names
        resources["laws_ready"] = COLLECTION_LAWS in collection_names

        if resources["contracts_ready"]:
            resources["embed_model"] = SentenceTransformer(EMBED_MODEL_CONTRACTS)
        if resources["laws_ready"]:
            resources["embed_model_laws"] = SentenceTransformer(EMBED_MODEL_LAWS)

    except Exception as e:
        print(f"❌ خطأ في تحميل قاعدة البيانات: {e}")

    return resources

# ──────────────────────────────────────────────────────────────────
# استخراج نص PDF استرشادي (اختياري)
# ──────────────────────────────────────────────────────────────────
def extract_pdf_text(pdf_path: str) -> str:
    text_content = ""
    try:
        reader = PdfReader(pdf_path)
        for page in reader.pages:
            text = page.extract_text()
            if text:
                text_content += text + "\n"
    except Exception as e:
        print(f"❌ خطأ في قراءة ملف الـ PDF: {e}")
    return text_content

# ──────────────────────────────────────────────────────────────────
# تحميل ملف أنواع العقود والبنود
# ──────────────────────────────────────────────────────────────────
def load_clause_types(json_path: str = CLAUSE_TYPES_PATH) -> dict:
    try:
        with open(json_path, "r", encoding="utf-8") as f:
            return json.load(f)
    except Exception as e:
        print(f"❌ خطأ في تحميل ملف أنواع العقود: {e}")
        return {}

def detect_contract_type(user_query: str, clause_types_db: dict | None = None) -> str | None:
    normalized_query = user_query.strip()
    best_match = None
    best_match_len = 0
    for type_key, keywords in CONTRACT_TYPE_KEYWORDS.items():
        if clause_types_db and type_key not in clause_types_db:
            continue
        for kw in keywords:
            # نستخدم فحص يدوي بدل \b لأن \b مش بيتعامل مع الحروف العربية صح
            # بنفحص إن الكلمة ليست جزء من كلمة أكبر
            kw_len = len(kw)
            start = 0
            found = False
            while True:
                idx = normalized_query.find(kw, start)
                if idx == -1:
                    break
                # فحص الحد الأيسر: لازم يكون بداية النص أو حرف غير عربي
                ok_left = (idx == 0) or not _is_arabic(normalized_query[idx - 1])
                # فحص الحد الأيمن: لازم يكون نهاية النص أو حرف غير عربي
                end_pos = idx + kw_len
                ok_right = (end_pos >= len(normalized_query)) or not _is_arabic(normalized_query[end_pos])
                if ok_left and ok_right:
                    found = True
                    break
                start = idx + 1
            if found and kw_len > best_match_len:
                best_match = type_key
                best_match_len = kw_len
    return best_match


def _is_arabic(ch: str) -> bool:
    """فحص هل الحرف ضمن نطاق الحروف العربية أو الهمزات أو التشكيل"""
    cp = ord(ch)
    return (
        (0x0621 <= cp <= 0x063A) or   # حروف عربية أساسية
        (0x0641 <= cp <= 0x0652) or   # حروف عربية إضافية + تشكيل
        (0x0671 <= cp <= 0x06D3) or   # أحرف موسعة
        (0x06D5 <= cp <= 0x06DC) or   # أحرف أخرى
        (0x06DE <= cp <= 0x06E8) or   # أحرف أخرى
        (0x06EA <= cp <= 0x06EF) or   # أحرف أخرى
        (0x06F0 <= cp <= 0x06F9)      # أرقام عربية
    )

# ──────────────────────────────────────────────────────────────────
# استرجاع سياق قانوني (قوانين ملزمة بس)
# ──────────────────────────────────────────────────────────────────
def _search_collection(query: str, qdrant_client, embed_model, collection_name: str,
                        top_k: int = 6, use_passage_prefix: bool = True,
                        score_threshold: float | None = None) -> list:
    try:
        search_query = f"query: {query}" if use_passage_prefix else query
        query_vector = embed_model.encode(search_query).tolist()

        results = qdrant_client.query_points(
            collection_name=collection_name,
            query=query_vector,
            limit=top_k,
            with_payload=["text", "page_content", "legal_domain", "article_number", "type"]
        ).points

        if score_threshold is not None:
            results = [p for p in results if (p.score or 0) >= score_threshold]

        return results
    except Exception as e:
        print(f"⚠️ خطأ في استرجاع السياق من {collection_name}: {e}")
        return []

def retrieve_laws_context(query: str, resources: dict, top_k: int = 6) -> str:
    qdrant_client = resources.get("qdrant_client")
    if not qdrant_client or not resources.get("laws_ready") or not resources.get("embed_model_laws"):
        return ""

    context_parts = []
    laws_points = _search_collection(
        query, qdrant_client, resources["embed_model_laws"], COLLECTION_LAWS,
        top_k=top_k, use_passage_prefix=False, score_threshold=LAWS_SCORE_THRESHOLD
    )
    for point in laws_points:
        payload = point.payload or {}
        text = payload.get("text") or payload.get("page_content") or payload.get("content", "")
        if text and len(text.strip()) >= 30:
            context_parts.append(f"📌 [مرجع قانوني ملزم] {text.strip()}")

    return "\n\n".join(context_parts)

# ──────────────────────────────────────────────────────────────────
# شبكة أمان: شيل أي رموز Markdown لو الموديل سرّبها
# ──────────────────────────────────────────────────────────────────
def strip_markdown_formatting(text: str) -> str:
    if not text:
        return text
    text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
    text = re.sub(r'__(.+?)__', r'\1', text)
    text = re.sub(r'(?<!\w)\*(.+?)\*(?!\w)', r'\1', text)
    text = re.sub(r'(?<!\w)_(.+?)_(?!\w)', r'\1', text)
    text = re.sub(r'^#{1,6}\s*', '', text, flags=re.MULTILINE)
    return text.strip()

# ──────────────────────────────────────────────────────────────────
# ✏️ تعديل 2 (جديد بالكامل): شبكة أمان ضد الـ degeneration loop
# (زي حالة بند 11 "أو تطويرها أو إنتاجها" اللي كررت لا نهائياً).
# مستقلة تماماً عن sampling params — بتقطع أي جملة/كلمة اتكررت
# 4 مرات متتالية أو أكتر وبتوقف النص عندها.
# ──────────────────────────────────────────────────────────────────
def truncate_repetition_loop(text: str, min_repeats: int = 4) -> str:
    if not text:
        return text
    words = text.split()
    for window in range(1, 6):  # نافذة من كلمة لغاية 5 كلمات
        i = 0
        while i + window * min_repeats <= len(words):
            chunk = words[i:i + window]
            repeats = 1
            j = i + window
            while j + window <= len(words) and words[j:j + window] == chunk:
                repeats += 1
                j += window
            if repeats >= min_repeats:
                return " ".join(words[:i + window]).strip()
            i += 1
    return text

# ──────────────────────────────────────────────────────────────────
# التوليد عبر Ollama
#
# ✏️ تعديل 2 (تكملة): ضفنا frequency_penalty + repeat_penalty (جوه
# extra_body بصياغة Ollama الأصلية كـ fallback لو الـ OpenAI-style
# param مش بيتطبق فعلياً عن طريق endpoint التوافق). ده أول خط دفاع
# قبل ما نوصل لـ truncate_repetition_loop.
# ──────────────────────────────────────────────────────────────────
def generate_with_ollama(messages: list, max_tokens: int) -> str:
    try:
        client = OpenAI(api_key="ollama", base_url=OLLAMA_BASE_URL)
        response = client.chat.completions.create(
            model=OLLAMA_MODEL,
            messages=messages,
            temperature=0.15,
            top_p=0.95,
            max_tokens=max_tokens,
            frequency_penalty=0.4,                    # ✏️ جديد
            extra_body={"repeat_penalty": 1.3},        # ✏️ جديد (صياغة Ollama الأصلية)
        )
        raw = response.choices[0].message.content
        cleaned = re.sub(r'oi[\s\S]*?sa', '', raw, flags=re.DOTALL)
        return cleaned.strip()
    except Exception as e:
        return (
            f"❌ خطأ في التوليد. تأكدي إن (1) أمر `ollama serve` شغال و"
            f"(2) المودل متحمّل بـ `ollama pull {OLLAMA_MODEL}`.\n\nتفاصيل الخطأ: {str(e)}"
        )

# ──────────────────────────────────────────────────────────────────
# صياغة بند واحد بالظبط
#
# ✏️ تعديل 1 (تكملة): شلنا legal_basis خالص (كان dead field أصلاً —
# الـ JSON بيرجّعه null دايماً دلوقتي ومكانش حتى بيتحط في الـ template).
# وغيّرنا تسمية laws_context_block من "استرشادي/للاستئناس فقط" لـ
# "المرجع القانوني (المصدر الملزم للأرقام)" عشان تتماشى مع الـ system
# prompt الجديد. وضفنا truncate_repetition_loop بعد التنضيف من الـ Markdown.
# ──────────────────────────────────────────────────────────────────
def generate_single_clause(clause: dict, contract_type_ar: str, laws_context: str) -> str:
    laws_context_block = (
        f"المرجع القانوني (المصدر الملزم للأرقام):\n{laws_context}\n"
        if laws_context else
        "المرجع القانوني: غير متاح لهذا البند — اكتبي صياغة عامة بدون أرقام.\n"
    )

    system_content = CLAUSE_SYSTEM_PROMPT.format(
        contract_type_ar=contract_type_ar,
        clause_title=clause.get("title", ""),
        clause_description=clause.get("description") or "بدون وصف إضافي",
        laws_context_block=laws_context_block,
    )
    messages = [
        {"role": "system", "content": system_content},
        {"role": "user", "content": "اكتبي نص هذا البند الآن."},
    ]
    raw = generate_with_ollama(messages, CLAUSE_MAX_TOKENS)
    cleaned = strip_markdown_formatting(raw)
    cleaned = truncate_repetition_loop(cleaned)  # ✏️ جديد

    # شبكة أمان: لو الموديل رجع بفارغ (أو التكرار قطع النص لحد حتة صغيرة أوي)،
    # نرجع وصف الـ JSON الأصلي كنص بديل
    if not cleaned or len(cleaned) < 20:
        fallback_text = clause.get("description") or "تم الاتفاق على هذا البند وفقاً للقانون."
        print(f"   ⚠️ تحذير: الموديل لم يُنتج نصاً صالحاً للبند [{clause.get('title')}]. تم استخدام النص الاحتياطي.")
        return fallback_text

    return cleaned

# ──────────────────────────────────────────────────────────────────
# بناء الديباجة (تمهيد + بيانات الطرفين) — قالب ثابت في بايثون
# ──────────────────────────────────────────────────────────────────
def build_static_preamble(contract_type_ar: str) -> str:
    return (
        f"عقد {contract_type_ar}\n\n"
        "إنه في يوم ........... الموافق ___/___/_______ م\n"
        "تحرر هذا العقد بين كل من:\n\n"
        "أولاً: الطرف الأول\n"
        "الاسم: [الاسم الكامل] | الرقم القومي: [14 رقم] | العنوان: [العنوان الكامل] | الصفة: [صفة الطرف الأول]\n\n"
        "ثانياً: الطرف الثاني\n"
        "الاسم: [الاسم الكامل] | الرقم القومي: [14 رقم] | العنوان: [العنوان الكامل] | الصفة: [صفة الطرف الثاني]\n\n"
        "تمهيد:\n"
        "بعد أن أقر الطرفان بأهليتهما القانونية الكاملة للتعاقد والتصرف، اتفقا على ما يلي:\n"
    )

def build_static_closing(contract_data: dict) -> str:
    lines = [
        "\nالطرف الأول: التوقيع ....................",
        "الطرف الثاني: التوقيع ....................\n",
        "المستندات القانونية الملحقة والمطلوبة للتوثيق:",
    ]
    notes = contract_data.get("notes") or []
    if notes:
        for note in notes:
            lines.append(f"• {note}")
    else:
        lines.append("• صورة من الهوية الوطنية للطرفين")
    return "\n".join(lines)

# ──────────────────────────────────────────────────────────────────
# فحص تأكيدي بسيط
# ──────────────────────────────────────────────────────────────────
def build_clause_validation_report(expected_count: int, generated_count: int) -> dict:
    return {
        "checked": True,
        "expected_count": expected_count,
        "found_count": generated_count,
        "missing_titles": [],
        "fuzzy_titles": [],
        "is_complete": expected_count == generated_count,
    }

# ──────────────────────────────────────────────────────────────────
# توليد ملف Word (.docx)
# ──────────────────────────────────────────────────────────────────
def create_word_document(contract_text: str, output_path: str = "عقد.docx") -> str:
    doc = Document()

    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    section = doc.sections[0]
    section.header_distance = Inches(0.5)

    lines = contract_text.split('\n')
    for line in lines:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p.paragraph_format.line_spacing = 1.3

        run = p.add_run(line)
        run.font.name = 'Arial'
        run.font.size = Pt(14)
        run.font.color.rgb = RGBColor(0, 0, 0)

    doc.save(output_path)
    return output_path

# ──────────────────────────────────────────────────────────────────
# دالة تجميعية: من سؤال المستخدم لحد ملف Word جاهز
#
# ✏️ تعديل 3: بقينا نستخدم search_keywords فعلياً في بناء استعلام
# الـ RAG لكل بند (كان قبل كده بيتبنى من title + description بس،
# والـ search_keywords كانت متضافة في الـ JSON من غير ما حد يقراها).
# ──────────────────────────────────────────────────────────────────
def generate_contract(user_query: str, pdf_context: str = "") -> dict:
    clause_types_db = load_clause_types()
    contract_type_key = detect_contract_type(user_query, clause_types_db)

    if not contract_type_key:
        print("⚠️ لم يتم تحديد نوع العقد من كلام المستخدم.")
        return {"contract_text": None, "docx_path": None, "contract_type_key": None, "clause_validation": {"checked": False}}

    contract_data = clause_types_db.get(contract_type_key, {})
    contract_type_ar = contract_data.get("contract_type_ar", "")
    specific_clauses = contract_data.get("specific_clauses", [])
    expected_n = len(specific_clauses)
    print(f"ℹ️ هيتم توليد {expected_n} بند لنوع العقد '{contract_type_key}' (مع RAG خاص لكل بند).")

    # تحميل موارد RAG مرة واحدة فقط للأداء
    rag_resources = load_rag_resources()

    body_parts = []
    for idx, clause in enumerate(specific_clauses, start=1):
        title = clause.get("title", "")
        description = clause.get("description") or ""
        keywords = clause.get("search_keywords") or []  # ✏️ جديد

        # ✏️ معدّل: استعلام الـ RAG بقى بيضم الـ search_keywords كمان،
        # مش بس العنوان والوصف — ده بيقوّي recall المصطلح القانوني الحرفي
        clause_search_query = f"{title}. {description}. {' '.join(keywords)}".strip()
        laws_context = retrieve_laws_context(clause_search_query, rag_resources, top_k=3)

        if pdf_context:
            laws_context = (laws_context + "\n\n" + pdf_context).strip()

        # توليد البند بالسياق الدقيق الخاص به
        clause_body = generate_single_clause(clause, contract_type_ar, laws_context)

        body_parts.append(f"البند {idx} — {title}: {clause_body}")
        print(f"   ✔️ اتولّد بند {idx}/{expected_n}: {title}")

    full_text = "\n\n".join(
        [build_static_preamble(contract_type_ar)] + body_parts + [build_static_closing(contract_data)]
    )

    clause_validation = build_clause_validation_report(expected_n, len(body_parts))
    docx_path = create_word_document(full_text)

    return {
        "contract_text": full_text,
        "docx_path": docx_path,
        "contract_type_key": contract_type_key,
        "clause_validation": clause_validation,
    }

# ──────────────────────────────────────────────────────────────────
# تشغيل مباشر من الـ terminal للتجربة
# ──────────────────────────────────────────────────────────────────
if __name__ == "__main__":
    query = input("📝 اكتبي طلب العقد: ")
    result = generate_contract(query)

    print("\n" + "=" * 60)
    print(f"🏷️ نوع العقد المكتشف: {result['contract_type_key'] or 'غير محدد'}")
    print("=" * 60)
    if result["contract_text"]:
        print(result["contract_text"])
    print("=" * 60)

    if result["docx_path"]:
        print(f"\n✔️ تم حفظ العقد في: {result['docx_path']}")
    else:
        print("\n⚠️ لم يتم استخراج نص عقد منظم من الرد.")

    cv = result.get("clause_validation") or {}
    if cv.get("checked"):
        status = "✅ مكتمل" if cv.get("is_complete") else "⚠️ ناقص"
        print(f"\n📋 تغطية البنود: {status} ({cv.get('found_count')}/{cv.get('expected_count')})")
